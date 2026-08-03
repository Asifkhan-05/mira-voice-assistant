import webbrowser
import os
import subprocess
import datetime
import requests
import sys
sys.path.append(r"D:\VoiceAssistant")
from memory import remember, recall, forget

# ─── App Paths ───────────────────────────────────────────
CHROME_PATH = r"C:\Program Files\Google\Chrome\Application\chrome.exe"
BRAVE_PATH = r"C:\Program Files\BraveSoftware\Brave-Browser\Application\brave.exe"
VSCODE_PATH = r"C:\Users\Asif\AppData\Local\Programs\Microsoft VS Code\Code.exe"
SPOTIFY_PATH = r"C:\Users\Asif\AppData\Roaming\Spotify\Spotify.exe"
STEAM_PATH = r"C:\Program Files (x86)\Steam\Steam.exe"
WORD_PATH = r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE"

WEATHER_API_KEY = "3873c793fa5cf1c8e21eabc34e171eef"

# ─── Browser Actions ─────────────────────────────────────
def open_chrome():
    if os.path.exists(CHROME_PATH):
        subprocess.Popen(CHROME_PATH)
        return "Opening Chrome!"
    return "Couldn't find Chrome on your system."

def open_brave():
    if os.path.exists(BRAVE_PATH):
        subprocess.Popen(BRAVE_PATH)
        return "Opening Brave!"
    return "Couldn't find Brave on your system."

def open_youtube():
    webbrowser.open("https://www.youtube.com")
    return "Opening YouTube!"

def open_google():
    webbrowser.open("https://www.google.com")
    return "Opening Google!"

def open_github():
    webbrowser.open("https://www.github.com")
    return "Opening GitHub!"

def open_gmail():
    webbrowser.open("https://mail.google.com")
    return "Opening Gmail!"

def open_maps():
    webbrowser.open("https://maps.google.com")
    return "Opening Google Maps!"

def open_translate():
    webbrowser.open("https://translate.google.com")
    return "Opening Google Translate!"

def search_youtube(query):
    url = f"https://www.youtube.com/results?search_query={query.replace(' ', '+')}"
    webbrowser.open(url)
    return f"Searching YouTube for {query}!"

def search_google(query):
    url = f"https://www.google.com/search?q={query.replace(' ', '+')}"
    webbrowser.open(url)
    return f"Searching Google for {query}!"

# ─── System Apps ─────────────────────────────────────────
def open_notepad():
    subprocess.Popen("notepad.exe")
    return "Opening Notepad!"

def open_calculator():
    subprocess.Popen("calc.exe")
    return "Opening Calculator!"

def open_word():
    if os.path.exists(WORD_PATH):
        subprocess.Popen(WORD_PATH)
        return "Opening Word!"
    return "I couldn't find Word on your system."

def open_excel():
    subprocess.Popen("excel.exe")
    return "Opening Excel!"

def open_powerpoint():
    subprocess.Popen("powerpnt.exe")
    return "Opening PowerPoint!"

def open_paint():
    subprocess.Popen("mspaint.exe")
    return "Opening Paint!"

def open_file_explorer():
    subprocess.Popen("explorer.exe")
    return "Opening File Explorer!"

def open_task_manager():
    subprocess.Popen("taskmgr.exe")
    return "Opening Task Manager!"

def open_settings():
    subprocess.Popen("ms-settings:", shell=True)
    return "Opening Settings!"

def open_cmd():
    subprocess.Popen("cmd.exe")
    return "Opening Command Prompt!"

def open_powershell():
    subprocess.Popen("powershell.exe")
    return "Opening PowerShell!"

def open_snipping_tool():
    subprocess.Popen("snippingtool.exe")
    return "Opening Snipping Tool!"

# ─── Dev Tools ───────────────────────────────────────────
def open_vscode():
    if os.path.exists(VSCODE_PATH):
        subprocess.Popen(VSCODE_PATH)
        return "Opening VS Code!"
    subprocess.Popen("code", shell=True)
    return "Opening VS Code!"

# ─── Entertainment ───────────────────────────────────────
def open_spotify():
    if os.path.exists(SPOTIFY_PATH):
        subprocess.Popen(SPOTIFY_PATH)
        return "Opening Spotify!"
    webbrowser.open("https://open.spotify.com")
    return "Opening Spotify in browser!"

def open_steam():
    if os.path.exists(STEAM_PATH):
        subprocess.Popen(STEAM_PATH)
        return "Opening Steam!"
    return "Couldn't find Steam on your system."

def open_netflix():
    webbrowser.open("https://www.netflix.com")
    return "Opening Netflix!"

def open_prime():
    webbrowser.open("https://www.primevideo.com")
    return "Opening Prime Video!"

# ─── System Controls ─────────────────────────────────────
def get_time():
    now = datetime.datetime.now()
    return f"It's {now.strftime('%I:%M %p')}."

def get_date():
    now = datetime.datetime.now()
    return f"Today is {now.strftime('%A, %B %d %Y')}."

def shutdown_pc():
    os.system("shutdown /s /t 5")
    return "Shutting down in 5 seconds!"

def restart_pc():
    os.system("shutdown /r /t 5")
    return "Restarting in 5 seconds!"

def sleep_pc():
    os.system("rundll32.exe powrprof.dll,SetSuspendState 0,1,0")
    return "Going to sleep!"

def lock_pc():
    os.system("rundll32.exe user32.dll,LockWorkStation")
    return "Locking your PC!"

def increase_volume():
    from ctypes import cast, POINTER
    from comtypes import CLSCTX_ALL
    from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume
    devices = AudioUtilities.GetSpeakers()
    interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
    volume = cast(interface, POINTER(IAudioEndpointVolume))
    current = volume.GetMasterVolumeLevelScalar()
    volume.SetMasterVolumeLevelScalar(min(1.0, current + 0.1), None)
    return "Volume increased!"

def decrease_volume():
    from ctypes import cast, POINTER
    from comtypes import CLSCTX_ALL
    from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume
    devices = AudioUtilities.GetSpeakers()
    interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
    volume = cast(interface, POINTER(IAudioEndpointVolume))
    current = volume.GetMasterVolumeLevelScalar()
    volume.SetMasterVolumeLevelScalar(max(0.0, current - 0.1), None)
    return "Volume decreased!"

def mute_volume():
    os.system("nircmd.exe mutesysvolume 1")
    return "Muted!"

# ─── Weather ─────────────────────────────────────────────
def get_weather(city="Bengaluru"):
    try:
        resp = requests.get(
            "https://api.openweathermap.org/data/2.5/weather",
            params={
                "q": city,
                "appid": WEATHER_API_KEY,
                "units": "metric"
            },
            timeout=10
        )
        data = resp.json()

        if resp.status_code != 200:
            return "Couldn't fetch weather right now."

        temp = data["main"]["temp"]
        feels_like = data["main"]["feels_like"]
        description = data["weather"][0]["description"]
        humidity = data["main"]["humidity"]
        city_name = data["name"]

        return f"It's {temp}°C in {city_name}, feels like {feels_like}°C, {description}, humidity is {humidity}%."

    except Exception as e:
        print(f"❌ Weather error: {e}")
        return "Something went wrong fetching the weather."

import threading

def set_timer(seconds, label="Timer"):
    def timer_done():
        print(f"⏰ {label} done!")
        from tts import speak
        speak(f"Hey! Your {label} is done!")
    
    timer = threading.Timer(seconds, timer_done)
    timer.start()
    
    if seconds >= 60:
        minutes = seconds // 60
        return f"Got it! Timer set for {minutes} minute{'s' if minutes > 1 else ''}!"
    return f"Got it! Timer set for {seconds} seconds!"

def parse_timer(text):
    """Extract duration from text like 'set a timer for 5 minutes'"""
    import re
    text_lower = text.lower()
    
    # Match patterns like "5 minutes", "30 seconds", "1 hour"
    hours = re.search(r'(\d+)\s*hour', text_lower)
    minutes = re.search(r'(\d+)\s*min', text_lower)
    seconds = re.search(r'(\d+)\s*sec', text_lower)
    
    total_seconds = 0
    label_parts = []
    
    if hours:
        total_seconds += int(hours.group(1)) * 3600
        label_parts.append(f"{hours.group(1)} hour")
    if minutes:
        total_seconds += int(minutes.group(1)) * 60
        label_parts.append(f"{minutes.group(1)} minute")
    if seconds:
        total_seconds += int(seconds.group(1))
        label_parts.append(f"{seconds.group(1)} second")
    
    if total_seconds > 0:
        label = " ".join(label_parts)
        return set_timer(total_seconds, label)
    
    return "I didn't catch the duration. Try saying set a timer for 5 minutes."
# ─── Search and Save to Word ─────────────────────────────
def search_and_save_to_word(query):
    from docx import Document
    import re

    print(f"🔍 Searching for: {query}")

    headers = {
        "User-Agent": "Athena/1.0 (Voice Assistant; contact@example.com)"
    }

    try:
        search_resp = requests.get(
            "https://en.wikipedia.org/w/api.php",
            params={
                "action": "query",
                "list": "search",
                "format": "json",
                "srsearch": query,
                "srlimit": 1
            },
            headers=headers,
            timeout=10
        )

        search_data = search_resp.json()
        results = search_data.get("query", {}).get("search", [])

        if not results:
            return f"Couldn't find anything about {query}."

        page_title = results[0]["title"]
        print(f"📄 Found page: {page_title}")

        content_resp = requests.get(
            "https://en.wikipedia.org/w/api.php",
            params={
                "action": "query",
                "format": "json",
                "titles": page_title,
                "prop": "extracts",
                "explaintext": True,
                "exsectionformat": "plain"
            },
            headers=headers,
            timeout=10
        )

        content_data = content_resp.json()
        pages = content_data.get("query", {}).get("pages", {})
        page = next(iter(pages.values()))
        content = page.get("extract", "")

        if not content:
            return "Found the page but couldn't extract content."

        doc = Document()
        doc.add_heading(page_title, 0)
        doc.add_paragraph(f"Source: https://en.wikipedia.org/wiki/{page_title.replace(' ', '_')}")
        doc.add_paragraph("")

        paragraphs = [p.strip() for p in content.split("\n") if p.strip()]
        for para in paragraphs[:80]:
            doc.add_paragraph(para)

        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        filename = re.sub(r'[^\w\s]', '', query)[:30].strip().replace(" ", "_")
        filepath = os.path.join(desktop, f"{filename}.docx")
        doc.save(filepath)

        print(f"💾 Saved to: {filepath}")
        return f"Done! Saved content about {page_title} to your desktop!"

    except Exception as e:
        print(f"❌ Error: {e}")
        return "Something went wrong while searching and saving."

# ─── Memory ──────────────────────────────────────────────
def handle_memory(text):
    text_lower = text.lower()

    if "my name is" in text_lower:
        after = text_lower.split("my name is")[-1].strip()
        name = after.split()[0].strip(".,!?").capitalize()
        remember("user_name", name)
        return f"Got it! I'll remember your name is {name}."

    if "i live in" in text_lower:
        after = text_lower.split("i live in")[-1].strip()
        city = after.split()[0].strip(".,!?").capitalize()
        remember("user_city", city)
        return f"Got it! I'll remember you live in {city}."

    if "i am from" in text_lower:
        after = text_lower.split("i am from")[-1].strip()
        place = after.split()[0].strip(".,!?").capitalize()
        remember("user_city", place)
        return f"Got it! I'll remember you're from {place}."

    if "i like" in text_lower:
        interest = text_lower.split("i like")[-1].strip().strip(".,!?")
        remember("user_interest", interest)
        return f"Got it! I'll remember you like {interest}."

    if "i love" in text_lower:
        interest = text_lower.split("i love")[-1].strip().strip(".,!?")
        remember("user_interest", interest)
        return f"Nice! I'll remember you love {interest}."

    if "forget my" in text_lower:
        key = text_lower.split("forget my")[-1].strip().strip(".,!?")
        forget(f"user_{key}")
        return f"Done, forgot your {key}."

    return None

# ─── Smart Keyword Intent Map ─────────────────────────────
INTENT_KEYWORDS = {
    "open_youtube":       ["youtube"],
    "open_chrome":        ["chrome"],
    "open_brave":         ["brave"],
    "open_google":        ["google"],
    "open_github":        ["github"],
    "open_gmail":         ["gmail", "mail", "email"],
    "open_maps":          ["maps", "google maps", "directions"],
    "open_translate":     ["translate", "translation"],
    "open_notepad":       ["notepad"],
    "open_calculator":    ["calculator", "calc"],
    "open_word":          ["word", "word document", "ms word", "microsoft word"],
    "open_excel":         ["excel", "spreadsheet"],
    "open_powerpoint":    ["powerpoint", "presentation", "slides"],
    "open_paint":         ["paint"],
    "open_file_explorer": ["file explorer", "explorer", "files", "folder"],
    "open_task_manager":  ["task manager", "tasks"],
    "open_settings":      ["settings", "setting"],
    "open_cmd":           ["command prompt", "cmd"],
    "open_powershell":    ["powershell"],
    "open_snipping_tool": ["snipping", "screenshot", "snip"],
    "open_vscode":        ["vs code", "vscode", "visual studio code", "code editor"],
    "open_spotify":       ["spotify", "music"],
    "open_steam":         ["steam", "gaming"],
    "open_netflix":       ["netflix"],
    "open_prime":         ["prime video", "prime", "amazon"],
    "get_time":           ["time", "what time"],
    "get_date":           ["date", "what date", "today"],
    "get_weather":        ["weather", "temperature", "how hot", "how cold"],
    "shutdown_pc":        ["shut down", "shutdown", "turn off"],
    "restart_pc":         ["restart", "reboot"],
    "sleep_pc":           ["sleep", "hibernate"],
    "lock_pc":            ["lock"],
    "set_timer": ["timer", "set a timer", "remind me in", "alarm"],
    "increase_volume":    ["volume up", "increase volume", "louder", "turn up"],
    "decrease_volume":    ["volume down", "decrease volume", "quieter", "turn down"],
    "mute_volume":        ["mute", "silence"],
}

INTENT_FUNCTIONS = {
    "open_youtube":       open_youtube,
    "open_chrome":        open_chrome,
    "open_brave":         open_brave,
    "open_google":        open_google,
    "open_github":        open_github,
    "open_gmail":         open_gmail,
    "open_maps":          open_maps,
    "open_translate":     open_translate,
    "open_notepad":       open_notepad,
    "open_calculator":    open_calculator,
    "open_word":          open_word,
    "open_excel":         open_excel,
    "open_powerpoint":    open_powerpoint,
    "open_paint":         open_paint,
    "open_file_explorer": open_file_explorer,
    "open_task_manager":  open_task_manager,
    "open_settings":      open_settings,
    "open_cmd":           open_cmd,
    "open_powershell":    open_powershell,
    "open_snipping_tool": open_snipping_tool,
    "open_vscode":        open_vscode,
    "open_spotify":       open_spotify,
    "open_steam":         open_steam,
    "open_netflix":       open_netflix,
    "open_prime":         open_prime,
    "get_time":           get_time,
    "get_date":           get_date,
    "get_weather":        get_weather,
    "shutdown_pc":        shutdown_pc,
    "restart_pc":         restart_pc,
    "sleep_pc":           sleep_pc,
    "lock_pc":            lock_pc,
    "increase_volume":    increase_volume,
    "decrease_volume":    decrease_volume,
    "mute_volume":        mute_volume,
}

# ─── Main Action Detector ─────────────────────────────────
def detect_action(text):
    text_lower = text.lower()

    # Memory first
    memory_result = handle_memory(text)
    if memory_result:
        return memory_result

    # Search and save to Word
    if any(phrase in text_lower for phrase in ["search and save", "find and save", "search and write", "look up and save"]):
        for kw in ["search and save", "find and save", "search and write", "look up and save"]:
            if kw in text_lower:
                query = text_lower.split(kw)[-1].strip().strip(".,!?")
                if query:
                    return search_and_save_to_word(query)

    # Weather with specific city
    if "weather" in text_lower or "temperature" in text_lower:
        for kw in ["weather in", "temperature in", "weather at"]:
            if kw in text_lower:
                city = text_lower.split(kw)[-1].strip().strip(".,!?")
                if city:
                    return get_weather(city)
        return get_weather()
        # Timer
    if any(w in text_lower for w in ["timer", "remind me in", "alarm"]):
         return parse_timer(text)


    # YouTube search
    if "youtube" in text_lower and any(w in text_lower for w in ["search", "find", "play", "look up"]):
        for kw in ["search for", "search", "find", "play", "look up"]:
            if kw in text_lower:
                query = text_lower.split(kw)[-1]
                query = query.replace("on youtube", "").replace("in youtube", "").replace("youtube", "").strip()
                if query:
                    return search_youtube(query)

    # Google search
    if any(w in text_lower for w in ["search", "find", "look up"]) and "youtube" not in text_lower:
        for kw in ["search for", "search", "find", "look up"]:
            if kw in text_lower:
                query = text_lower.split(kw)[-1]
                query = query.replace("on google", "").replace("in google", "").replace("google", "").strip()
                if query:
                    return search_google(query)

    # Match keywords
    sorted_intents = sorted(INTENT_KEYWORDS.items(), key=lambda x: max(len(k) for k in x[1]), reverse=True)
    for intent, keywords in sorted_intents:
        for keyword in keywords:
            if keyword in text_lower:
                print(f"🎯 Intent: {intent}")
                return INTENT_FUNCTIONS[intent]()

    return None